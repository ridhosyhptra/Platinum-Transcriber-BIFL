"""
Platinum Transcriber — desktop transcription application (CustomTkinter GUI).

Architecture (enterprise-oriented)
==================================
- **UI thread**: Tk/CTk main loop; never blocks on I/O or model inference.
- **Worker thread**: `logic()` runs transcription; communicates via `queue.Queue`
  (producer-consumer). UI updates use `msg_queue` + `poll_queue()` (100 ms tick).
- **Snapshot pattern**: `start_thread()` copies UI state into an immutable dict
  before spawning the worker — avoids races with Tk variables.
- **Streaming I/O**: Faster-Whisper segments are written to disk incrementally
  (bounded memory vs. loading full transcripts).
- **Optional live dictation**: Producer-consumer with PyAudio + Whisper/Google;
  `threading.Lock` protects `live_transcript`.

Data structures
===============
- `files_to_do`: list of (filesystem_path, output_basename) — O(n) batch queue.
- `msg_queue` / `live_audio_queue`: FIFO queues (thread-safe).
- `VAD_PRESETS`, i18n dicts: O(1) lookup by key.
- SQLite `logs`: append-only analytics (see `init_db` / `save_log`).

Security & ops
==============
- User prefs and API keys: `~/PlatinumConfig.json` — do not commit real keys.
- Hugging Face downloads: `HF_HUB_DISABLE_XET` avoids Windows Xet transport issues.

See `docs/DEVELOPER.md` for module map and extension points.
"""

from __future__ import annotations

import os
import sys
import time
import threading
import json
import subprocess
import urllib.request
import shutil
import queue
import sqlite3
import datetime
import re
import locale
import ctypes
from ctypes import wintypes
import math
import tkinter as tk
from tkinter import filedialog, messagebox
from typing import Optional

os.environ["HF_HUB_DISABLE_XET"] = "1"

try:
    import pyaudio
    HAS_PYAUDIO = True
except ImportError:
    HAS_PYAUDIO = False

try:
    import numpy as np_audio
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False

try:
    import speech_recognition as sr
    HAS_SPEECH_REC = True
except ImportError:
    HAS_SPEECH_REC = False

# --- 1. INJEKSI FFMPEG PORTABLE (100% PLUG AND PLAY) ---
def inject_ffmpeg():
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS 
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    
    ffmpeg_path = os.path.join(base_path, "ffmpeg.exe")
    if os.path.exists(ffmpeg_path):
        os.environ["PATH"] += os.pathsep + base_path
        return True
    return False

HAS_FFMPEG = inject_ffmpeg()

def get_windows_desktop_dir():
    """Best-effort: dapatkan path Desktop yang benar (termasuk OneDrive redirect)."""
    if os.name != "nt":
        return os.path.join(os.path.expanduser("~"), "Desktop")

    # Prefer env if present
    userprofile = os.environ.get("USERPROFILE")
    onedrive = os.environ.get("OneDrive")
    if onedrive:
        candidate = os.path.join(onedrive, "Desktop")
        if os.path.isdir(candidate):
            return candidate
    if userprofile:
        candidate = os.path.join(userprofile, "Desktop")
        if os.path.isdir(candidate):
            return candidate

    # Known Folder: Desktop
    FOLDERID_Desktop = ctypes.c_byte * 16
    fid = FOLDERID_Desktop(
        0xB4, 0xBF, 0xCC, 0x3A, 0xDB, 0x2C, 0x42, 0x4C,
        0xB0, 0x29, 0x7F, 0xE9, 0x9A, 0x87, 0xC6, 0x41
    )
    path_ptr = wintypes.LPWSTR()
    try:
        shell32 = ctypes.windll.shell32
        ole32 = ctypes.windll.ole32
        hr = shell32.SHGetKnownFolderPath(ctypes.byref(fid), 0, 0, ctypes.byref(path_ptr))
        if hr == 0 and path_ptr.value and os.path.isdir(path_ptr.value):
            return path_ptr.value
    except Exception:
        pass
    finally:
        try:
            if path_ptr:
                ctypes.windll.ole32.CoTaskMemFree(path_ptr)
        except Exception:
            pass

    return os.path.join(os.path.expanduser("~"), "Desktop")

# --- 2. PRE-FLIGHT CHECK ---
missing_libs = []
for lib in ['yt_dlp', 'faster_whisper', 'deep_translator', 'deepl', 'docx', 'reportlab', 'psutil', 'assemblyai', 'customtkinter']:
    try: __import__(lib)
    except Exception: missing_libs.append(lib)

if missing_libs:
    tk.Tk().withdraw()
    # Hindari getdefaultlocale() (deprecated di Python baru)
    sys_lang = ((locale.getlocale()[0] or os.environ.get("LANG") or os.environ.get("LC_ALL") or "")).lower()
    use_id = sys_lang.startswith("id")
    if use_id:
        title = "Diagnostik Sistem Gagal"
        err_msg = "Library belum terinstal:\n" + "\n".join(missing_libs) + "\n\nJalankan skrip build/installer untuk menginstal otomatis."
    else:
        title = "System Diagnostics Failed"
        err_msg = "Missing libraries:\n" + "\n".join(missing_libs) + "\n\nRun the build/installer script to install them automatically."
    messagebox.showerror(title, err_msg)
    sys.exit()

import yt_dlp
from faster_whisper import WhisperModel
from deep_translator import GoogleTranslator
import deepl
import docx
from reportlab.pdfgen import canvas
import psutil
import assemblyai as aai
import customtkinter as ctk

ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")
CONFIG_FILE = os.path.join(os.path.expanduser("~"), "PlatinumConfig.json")
DB_FILE = os.path.join(os.path.expanduser("~"), "PlatinumLogs.db")

# Tunable limits (avoid magic numbers scattered in UI/backend)
CONSOLE_LOG_MAX_LINES = 5000
CONSOLE_LOG_TRIM_TO = 4000
LARGE_FILE_WARN_GB = 2.0
MIN_FREE_RAM_GB = 1.5

# --- 3. ANALYTICS DATABASE (SQLITE) ---
def init_db() -> None:
    conn = sqlite3.connect(DB_FILE)
    try:
        conn.execute("CREATE TABLE IF NOT EXISTS logs (id INTEGER PRIMARY KEY, tanggal TEXT, file_name TEXT, durasi_detik REAL, engine TEXT)")
        conn.commit()
    finally:
        conn.close()

def save_log(file_name: str, durasi: float, engine: str) -> None:
    try:
        conn = sqlite3.connect(DB_FILE)
        try:
            tgl = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            conn.execute("INSERT INTO logs (tanggal, file_name, durasi_detik, engine) VALUES (?, ?, ?, ?)", (tgl, file_name, durasi, engine))
            conn.commit()
        finally:
            conn.close()
    except Exception as e:
        print(f"[DB Error] Gagal menyimpan log: {e}")

init_db()

# --- 4. MODULE-LEVEL UTILITY FUNCTIONS ---
def _wrap_lines(text, max_chars=90):
    words = text.split()
    lines, cur = [], ""
    for w in words:
        if cur and len(cur) + 1 + len(w) > max_chars:
            lines.append(cur)
            cur = w
        else:
            cur = f"{cur} {w}" if cur else w
    if cur:
        lines.append(cur)
    return lines or [""]

def _text_to_paragraphs(text, sentences_per_para=5):
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    paragraphs, buf = [], []
    for sent in sentences:
        buf.append(sent)
        if len(buf) >= sentences_per_para:
            paragraphs.append(" ".join(buf))
            buf = []
    if buf:
        paragraphs.append(" ".join(buf))
    return paragraphs or [text]

def _fmt_ms(ms):
    s = ms // 1000
    h, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"

def _fmt_ass_ts(seconds):
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds % 60
    return f"{h:d}:{m:02d}:{s:05.2f}"

def _ass_header():
    return (
        "[Script Info]\nTitle: Platinum Transcriber\nScriptType: v4.00+\n"
        "WrapStyle: 0\nPlayResX: 1920\nPlayResY: 1080\n\n"
        "[V4+ Styles]\nFormat: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, "
        "OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, "
        "Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding\n"
        "Style: Default,Arial,48,&H00FFFFFF,&H000000FF,&H00000000,&H64000000,"
        "-1,0,0,0,100,100,0,0,1,2,1,2,10,10,30,1\n\n"
        "[Events]\nFormat: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text\n"
    )

def _fmt_sbv_ts(seconds):
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds % 60
    ms = int((s - int(s)) * 1000)
    return f"{h:d}:{m:02d}:{int(s):02d}.{ms:03d}"

def _ensure_ffmpeg_available():
    if HAS_FFMPEG:
        return True
    try:
        return subprocess.run(["ffmpeg", "-version"], capture_output=True).returncode == 0
    except FileNotFoundError:
        return False

VAD_PRESETS = {
    "low":  {"threshold": 0.35, "min_silence_duration_ms": 2500, "speech_pad_ms": 600},
    "med":  {"threshold": 0.45, "min_silence_duration_ms": 1500, "speech_pad_ms": 500},
    "high": {"threshold": 0.60, "min_silence_duration_ms": 800,  "speech_pad_ms": 300},
}

# --- LIVE DICTATION CONSTANTS & UTILITIES ---
LIVE_RATE = 16000
LIVE_CHUNK = 1024
LIVE_RECORD_SECONDS = 3
LIVE_SILENCE_THRESH = 0.01

def _rms_energy(audio_f32):
    if not HAS_NUMPY or len(audio_f32) == 0:
        return 0.0
    return math.sqrt(np_audio.mean(audio_f32 ** 2))

def _render_docx(filepath, title, engine_label, lines):
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = docx.shared.Pt(11)
    doc.add_heading(title, level=1)
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Transcribed: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Engine: {engine_label}")
    meta_run.italic = True
    meta_run.font.size = docx.shared.Pt(9)
    meta_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
    doc.add_paragraph("")
    for line in lines:
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(filepath)

def _render_pdf(filepath, title, engine_label, lines):
    c = canvas.Canvas(filepath)
    width, margin_x, top_y, line_h = 595, 40, 780, 14
    usable = width - 2 * margin_x
    max_chars = int(usable / 5.5)
    y = top_y
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin_x, y, title)
    y -= 20
    c.setFont("Helvetica", 8)
    c.setFillColorRGB(0.5, 0.5, 0.5)
    c.drawString(margin_x, y, f"Transcribed: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Engine: {engine_label}")
    y -= 10
    c.setStrokeColorRGB(0.8, 0.8, 0.8)
    c.line(margin_x, y, width - margin_x, y)
    y -= 18
    c.setFillColorRGB(0, 0, 0)
    c.setFont("Helvetica", 10)
    for line in lines:
        wrapped = _wrap_lines(line, max_chars) if line.strip() else [""]
        for wl in wrapped:
            if y < 40:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = top_y
            c.drawString(margin_x, y, wl)
            y -= line_h
    c.save()

# --- 5. STDOUT/STDERR -> UI QUEUE (recursion guard) ---
class QueueRedirector:
    """Redirects print() to the UI log queue; `_writing` prevents re-entrant recursion."""

    def __init__(self, q: queue.Queue) -> None:
        self.q = q
        self._writing = False

    def write(self, text):
        if text.strip() and not self._writing:
            self._writing = True
            try:
                self.q.put({"type": "log", "msg": text + "\n"})
            finally:
                self._writing = False

    def flush(self):
        pass

# --- 6. MAIN APPLICATION (CustomTkinter shell) ---
class PlatinumTranscriberApp(ctk.CTk):
    """
    Root window: tabs for input, output/translation, cloud options; console log;
    background worker (`logic`) for batch transcription.

    Threading contract: UI callbacks from worker threads must go through
    `msg_queue` (log/progress/ui) or `ui_call()`; never touch Tk widgets directly
    from `logic()` or live threads.
    """

    def __init__(self) -> None:
        super().__init__()
        self.title("Platinum Transcriber V8.0 (Enterprise Edition)")
        self.geometry("1100x850")
        self.minsize(950, 720)
        self.stop_flag = False
        self.msg_queue = queue.Queue()
        
        self.setup_vars()
        self.i18n = {
            "id": {
                "tab1": "1. Input & Mesin AI",
                "tab2": "2. Output & Terjemahan",
                "tab3": "3. AssemblyAI & Cloud",
                "btn_lang_to_en": "🔄 English",
                "btn_lang_to_id": "🔄 Indonesia",
                "btn_hist": "📊 RIWAYAT",
                "btn_diag": "🔍 DIAGNOSTIK SISTEM",
                "btn_start": "▶ MULAI EKSEKUSI",
                "btn_stop": "⏹ BATALKAN",
                "btn_stop_busy": "Membatalkan...",
                "sec_src": "Sumber Audio/Video",
                "src_yt": "Link YouTube (~30MB/jam)",
                "src_local": "File Lokal (Batch - 100% Offline)",
                "ph_yt": "Tempel link YouTube...",
                "browse": "Cari File",
                "no_file": "Belum ada file",
                "sec_model": "Ketajaman AI (Model Lokal)",
                "model_small": "Whisper SMALL (Sangat Cepat | RAM > 1.5GB)",
                "model_medium": "Whisper MEDIUM (Presisi Tinggi | RAM > 3.0GB)",
                "sum_exec": "Ringkasan Eksekusi",
                "sec_out": "Format Output Dokumen",
                "sec_trans": "Opsi Terjemahan (Otomatis ID/EN)",
                "tr_none": "Teks Original (100% Offline)",
                "tr_google": "Google Translate (EN->ID | Gratis)",
                "tr_any2en": "Google Translate (Any->EN | Gratis)",
                "tr_deepl": "DeepL AI (Presisi Tinggi EN->ID)",
                "ph_deepl": "DeepL API Key",
                "lbl_deepl_key": "DeepL API Key",
                "sum_result": "Ringkasan Hasil",
                "sec_pre": "Audio Pre-Processing (Lokal)",
                "pre_ffmpeg": "FFmpeg Noise Reduction (Pembersih Bising)",
                "sec_cloud": "AssemblyAI Enterprise Cloud",
                "use_aai": "Aktifkan AssemblyAI (Mematikan mesin lokal)",
                "aai_chap": "Auto Chapters",
                "aai_pii": "Redaksi PII",
                "aai_diar": "Diarization",
                "aai_filler": "Hapus Filler Word",
                "ph_vocab": "Custom Vocab (Koma: safety, hazmat, dll)",
                "ph_vocab_opt": "Custom Vocab (Opsional, pisahkan dengan koma)",
                "ph_aai_key": "AssemblyAI API Key",
                "lbl_aai_key": "AssemblyAI API Key  (Wajib)",
                "lbl_vocab": "Custom Vocabulary  (Opsional)",
                "sum_cloud": "Ringkasan Cloud",
                "hist_title": "Dashboard Analitik",
                "hist_lbl": "Riwayat Transkripsi (Log Database)",
                "sum_mode": "📌 Mode: {mode}\n📄 Target: {target}",
                "sum_engine": "🧠 Mesin AI: {engine}",
                "sum_out": "📁 Format Dihasilkan: {out}",
                "sum_tr": "🌐 Terjemahan: {tr}",
                "mode_yt": "Link YouTube",
                "mode_local": "File Lokal (Offline)",
                "target_files": "{n} File Dipilih",
                "file_count": "{n} file",
                "engine_small": "Whisper SMALL (Cepat)",
                "engine_medium": "Whisper MEDIUM (Presisi)",
                "engine_cloud": "AssemblyAI (Cloud) ⚠️",
                "tr_map_none": "Tidak Ada (Asli)",
                "tr_map_google": "Google Translate (EN→ID)",
                "tr_map_any2en": "Google Translate (Any→EN)",
                "tr_map_deepl": "DeepL Pro AI",
                "cloud_off": "Status: TIDAK AKTIF (100% Offline)\nLaptop Anda aman bekerja tanpa internet.",
                "cloud_on": "Status: AKTIF (Koneksi Cloud)\nFitur Ekstra: {opts}",
                "cloud_std": "Standar",
                "opt_chap": "Chapters",
                "opt_pii": "PII Redaction",
                "opt_diar": "Diarization",
                "mb_title_success": "Sukses",
                "mb_msg_success": "100% Selesai! Seluruh file telah diletakkan di Desktop.",
                "mb_title_cancelled": "Dibatalkan",
                "mb_msg_cancelled": "Soft E-Stop berhasil. File yang setengah jalan tetap tersimpan aman.",
                "mb_title_error": "System Error",
                "mb_title_diag": "Pre-Flight Check",
                "mb_msg_diag": "Internet Connection: {inet}\nFFmpeg Engine: {ff}\nRAM Free: {fr:.2f}GB / {tr:.2f}GB\nStorage Free: {fs:.2f}GB",
                "btn_live": "🎙️ REKAM & TRANSKIP (LIVE)",
                "btn_live_stop": "⏹ BERHENTI REKAM",
                "live_engine_google": "Google Speech (Online)",
                "live_engine_whisper": "Whisper Lokal (Offline)",
                "live_starting": "[LIVE] Memulai sesi live...",
                "live_warmup": "[LIVE] Memanaskan mesin Whisper (pertama kali mungkin lambat)...",
                "live_mic_ready": "[LIVE] Mikrofon aktif. Silakan berbicara...",
                "live_stopped": "[LIVE] Sesi live dihentikan.",
                "live_saved": "[LIVE] Transkrip disimpan di: {path}",
                "live_no_pyaudio": "Library 'pyaudio' belum terinstal.\nJalankan: pip install pyaudio",
                "live_no_speech": "Library 'SpeechRecognition' belum terinstal.\nJalankan: pip install SpeechRecognition",
                "live_no_numpy": "Library 'numpy' belum terinstal.\nJalankan: pip install numpy",
                "live_empty": "[LIVE] Tidak ada teks yang berhasil ditranskrip.",
                "warn_large_file": "[PERINGATAN] File '{name}' berukuran {size:.1f} GB. Ekstraksi audio awal akan memakan waktu ekstra. Harap bersabar.",
                "fmt_total": "[SISTEM] {n} file dipilih ({total_fmt}). Format terdeteksi: {exts}",
                "out_none": "Belum dipilih!",
                "use_vad": "VAD Filter (Lewati Bagian Diam/Sunyi)",
                "vad_level": "Sensitivitas VAD",
                "vad_low": "Rendah (Konservatif)",
                "vad_med": "Sedang (Direkomendasikan)",
                "vad_high": "Tinggi (Agresif)",
                "use_word_ts": "Word-Level Timestamps (Presisi Per-Kata)",
                "word_ts_warn": "Menambah waktu proses ~30%.",
            },
            "en": {
                "tab1": "1. Input & AI Engine",
                "tab2": "2. Output & Translation",
                "tab3": "3. AssemblyAI & Cloud",
                "btn_lang_to_en": "🔄 English",
                "btn_lang_to_id": "🔄 Indonesia",
                "btn_hist": "📊 HISTORY",
                "btn_diag": "🔍 SYSTEM DIAGNOSTICS",
                "btn_start": "▶ START EXECUTION",
                "btn_stop": "⏹ ABORT",
                "btn_stop_busy": "Aborting...",
                "sec_src": "Audio/Video Source",
                "src_yt": "YouTube Link (~30MB/hour)",
                "src_local": "Local Files (Batch - 100% Offline)",
                "ph_yt": "Paste YouTube link...",
                "browse": "Browse Files",
                "no_file": "No file selected",
                "sec_model": "Local AI Model",
                "model_small": "Whisper SMALL (Very Fast | Free RAM > 1.5GB)",
                "model_medium": "Whisper MEDIUM (Higher Accuracy | Free RAM > 3.0GB)",
                "sum_exec": "Execution Summary",
                "sec_out": "Output Formats",
                "sec_trans": "Translation Options",
                "tr_none": "Original Text (100% Offline)",
                "tr_google": "Google Translate (EN→ID | Free)",
                "tr_any2en": "Translate to English (Any→EN | Free)",
                "tr_deepl": "DeepL (High Accuracy EN→ID)",
                "ph_deepl": "DeepL API Key",
                "lbl_deepl_key": "DeepL API Key",
                "sum_result": "Output Summary",
                "sec_pre": "Audio Pre-Processing (Local)",
                "pre_ffmpeg": "FFmpeg Noise Reduction",
                "sec_cloud": "AssemblyAI Cloud",
                "use_aai": "Enable AssemblyAI (disables local engine)",
                "aai_chap": "Auto Chapters",
                "aai_pii": "PII Redaction",
                "aai_diar": "Diarization",
                "aai_filler": "Remove filler words",
                "ph_vocab": "Custom vocab (comma separated)",
                "ph_vocab_opt": "Custom vocab (optional, comma separated)",
                "ph_aai_key": "AssemblyAI API Key",
                "lbl_aai_key": "AssemblyAI API Key  (Required)",
                "lbl_vocab": "Custom Vocabulary  (Optional)",
                "sum_cloud": "Cloud Summary",
                "hist_title": "Analytics Dashboard",
                "hist_lbl": "Transcription History (DB Logs)",
                "sum_mode": "📌 Mode: {mode}\n📄 Target: {target}",
                "sum_engine": "🧠 Engine: {engine}",
                "sum_out": "📁 Outputs: {out}",
                "sum_tr": "🌐 Translation: {tr}",
                "mode_yt": "YouTube Link",
                "mode_local": "Local Files (Offline)",
                "target_files": "{n} file(s) selected",
                "file_count": "{n} file(s)",
                "engine_small": "Whisper SMALL (Fast)",
                "engine_medium": "Whisper MEDIUM (Accurate)",
                "engine_cloud": "AssemblyAI (Cloud) ⚠️",
                "tr_map_none": "None (Original)",
                "tr_map_google": "Google Translate (EN→ID)",
                "tr_map_any2en": "Translate to English (Any→EN)",
                "tr_map_deepl": "DeepL",
                "cloud_off": "Status: OFF (100% Offline)\nNo internet required.",
                "cloud_on": "Status: ON (Cloud)\nExtra features: {opts}",
                "cloud_std": "Standard",
                "opt_chap": "Chapters",
                "opt_pii": "PII Redaction",
                "opt_diar": "Diarization",
                "mb_title_success": "Success",
                "mb_msg_success": "Done! All output files have been saved to your Desktop.",
                "mb_title_cancelled": "Cancelled",
                "mb_msg_cancelled": "Soft E-Stop completed. Partially processed files were saved safely.",
                "mb_title_error": "System Error",
                "mb_title_diag": "Pre-Flight Check",
                "mb_msg_diag": "Internet Connection: {inet}\nFFmpeg Engine: {ff}\nRAM Free: {fr:.2f}GB / {tr:.2f}GB\nStorage Free: {fs:.2f}GB",
                "btn_live": "🎙️ RECORD & TRANSCRIBE (LIVE)",
                "btn_live_stop": "⏹ STOP RECORDING",
                "live_engine_google": "Google Speech (Online)",
                "live_engine_whisper": "Whisper Local (Offline)",
                "live_starting": "[LIVE] Starting live session...",
                "live_warmup": "[LIVE] Warming up Whisper engine (first time may be slow)...",
                "live_mic_ready": "[LIVE] Mic active. Start speaking...",
                "live_stopped": "[LIVE] Live session stopped.",
                "live_saved": "[LIVE] Transcript saved to: {path}",
                "live_no_pyaudio": "Library 'pyaudio' is not installed.\nRun: pip install pyaudio",
                "live_no_speech": "Library 'SpeechRecognition' is not installed.\nRun: pip install SpeechRecognition",
                "live_no_numpy": "Library 'numpy' is not installed.\nRun: pip install numpy",
                "live_empty": "[LIVE] No text was transcribed.",
                "warn_large_file": "[WARNING] File '{name}' is {size:.1f} GB. Initial audio extraction will take extra time. Please be patient.",
                "fmt_total": "[SYSTEM] {n} file(s) selected ({total_fmt}). Detected formats: {exts}",
                "out_none": "Not selected!",
                "use_vad": "VAD Filter (Skip Silent Parts)",
                "vad_level": "VAD Sensitivity",
                "vad_low": "Low (Conservative)",
                "vad_med": "Medium (Recommended)",
                "vad_high": "High (Aggressive)",
                "use_word_ts": "Word-Level Timestamps (Per-Word Precision)",
                "word_ts_warn": "Adds ~30% processing time.",
            },
        }
        self.current_lang = getattr(self, "current_lang", "id")
        if self.current_lang not in ("id", "en"):
            self.current_lang = "id"
        self.load_prefs()
        self.setup_ui()
        
        sys.stdout = QueueRedirector(self.msg_queue)
        sys.stderr = QueueRedirector(self.msg_queue)
        self.poll_queue()
        self.protocol("WM_DELETE_WINDOW", self._on_closing)

    def poll_queue(self) -> None:
        """Drain UI message queue (thread-safe consumer); schedules self recursively."""
        try:
            while True:
                try:
                    msg = self.msg_queue.get_nowait()
                except queue.Empty:
                    break
                if msg["type"] == "log":
                    self.console_log.configure(state="normal")
                    self.console_log.insert("end", msg["msg"])
                    line_count = int(self.console_log.index("end-1c").split(".")[0])
                    if line_count > CONSOLE_LOG_MAX_LINES:
                        self.console_log.delete("1.0", f"{line_count - CONSOLE_LOG_TRIM_TO}.0")
                    self.console_log.see("end")
                    self.console_log.configure(state="disabled")
                elif msg["type"] == "progress":
                    self.prog_bar.set(msg["val"])
                    self.prog_lbl.configure(text=f"{int(msg['val'] * 100)} %")
                elif msg["type"] == "prog_mode":
                    self.prog_bar.configure(mode=msg["val"])
                    if msg["val"] == "indeterminate": self.prog_bar.start()
                    else: self.prog_bar.stop()
                elif msg["type"] == "live_text":
                    self.console_log.configure(state="normal")
                    self.console_log.insert("end", msg["text"] + "\n")
                    line_count = int(self.console_log.index("end-1c").split(".")[0])
                    if line_count > CONSOLE_LOG_MAX_LINES:
                        self.console_log.delete("1.0", f"{line_count - CONSOLE_LOG_TRIM_TO}.0")
                    self.console_log.see("end")
                    self.console_log.configure(state="disabled")
                elif msg["type"] == "ui":
                    msg["fn"](*msg["args"], **msg["kwargs"])
        except Exception as e:
            try:
                print(f"[UI QUEUE ERROR] {e}")
            except Exception:
                pass
        self.after(100, self.poll_queue)

    def _on_closing(self):
        if self.live_recording:
            self.live_stop_event.set()
            self.live_recording = False
        self.stop_flag = True
        try:
            self.destroy()
        except Exception:
            pass

    def _reset_live_ui_state(self):
        t = self.i18n[self.current_lang]
        self.live_recording = False
        self.btn_live.configure(text=t["btn_live"], fg_color="#7B2D8E", hover_color="#5E1F6E")
        self.btn_start.configure(state="normal")
        self.live_engine_menu.configure(state="normal")
        self.btn_lang.configure(state="normal")
        self.btn_stop.configure(state="disabled")
        self.msg_queue.put({"type": "prog_mode", "val": "determinate"})
        self.msg_queue.put({"type": "progress", "val": 0})

    def ui_call(self, fn, *args, **kwargs):
        self.msg_queue.put({"type": "ui", "fn": fn, "args": args, "kwargs": kwargs})

    def open_folder(self, path: str):
        """UI-side helper: buka folder output di file explorer."""
        try:
            if os.name == "nt":
                os.startfile(path)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", path])
            else:
                subprocess.Popen(["xdg-open", path])
        except Exception as e:
            print(f"[OPEN FOLDER WARNING] {e}")

    def setup_vars(self):
        self.src_type = ctk.StringVar(value="local"); self.ai_eng = ctk.StringVar(value="small")
        self.out_srt = ctk.BooleanVar(value=True); self.out_vtt = ctk.BooleanVar(value=False)
        self.out_txt = ctk.BooleanVar(value=True); self.out_docx = ctk.BooleanVar(value=False); self.out_pdf = ctk.BooleanVar(value=False)
        self.out_json = ctk.BooleanVar(value=False); self.out_tsv = ctk.BooleanVar(value=False)
        self.out_sbv = ctk.BooleanVar(value=False); self.out_md = ctk.BooleanVar(value=False); self.out_ass = ctk.BooleanVar(value=False)
        self.trans_opt = ctk.StringVar(value="none"); self.clean_audio = ctk.BooleanVar(value=False)
        self.use_aai = ctk.BooleanVar(value=False)
        self.aai_chap = ctk.BooleanVar(value=False); self.aai_pii = ctk.BooleanVar(value=False)
        self.aai_diar = ctk.BooleanVar(value=False); self.aai_filler = ctk.BooleanVar(value=True)
        self.use_vad = ctk.BooleanVar(value=False); self.vad_level = ctk.StringVar(value="med")
        self.word_ts = ctk.BooleanVar(value=False)
        self.deepl_key = ctk.StringVar(); self.aai_key = ctk.StringVar(); self.aai_vocab = ctk.StringVar()
        self.url_input = ctk.StringVar(); self.local_files = []
        self.live_engine = ctk.StringVar(value="google")
        self.live_recording = False
        self.live_stop_event = threading.Event()
        self.live_audio_queue = queue.Queue()
        self.live_transcript = []
        self._live_lock = threading.Lock()

    def save_prefs(self):
        prefs = {
            "ai_eng": self.ai_eng.get(), "out_srt": self.out_srt.get(),
            "out_vtt": self.out_vtt.get(), "out_txt": self.out_txt.get(),
            "out_docx": self.out_docx.get(), "out_pdf": self.out_pdf.get(),
            "out_json": self.out_json.get(), "out_tsv": self.out_tsv.get(),
            "out_sbv": self.out_sbv.get(), "out_md": self.out_md.get(), "out_ass": self.out_ass.get(),
            "trans_opt": self.trans_opt.get(), "clean_audio": self.clean_audio.get(),
            "use_vad": self.use_vad.get(), "vad_level": self.vad_level.get(),
            "word_ts": self.word_ts.get(),
            "deepl_key": self.deepl_key.get(), "aai_key": self.aai_key.get(),
        }
        try:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(prefs, f)
        except Exception:
            pass

    def load_prefs(self):
        if not os.path.exists(CONFIG_FILE):
            return
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                p = json.load(f)
            self.ai_eng.set(p.get("ai_eng", "small"))
            self.trans_opt.set(p.get("trans_opt", "none"))
            self.out_srt.set(p.get("out_srt", True))
            self.out_vtt.set(p.get("out_vtt", False))
            self.out_txt.set(p.get("out_txt", True))
            self.out_docx.set(p.get("out_docx", False))
            self.out_pdf.set(p.get("out_pdf", False))
            self.out_json.set(p.get("out_json", False))
            self.out_tsv.set(p.get("out_tsv", False))
            self.out_sbv.set(p.get("out_sbv", False))
            self.out_md.set(p.get("out_md", False))
            self.out_ass.set(p.get("out_ass", False))
            self.clean_audio.set(p.get("clean_audio", False))
            self.use_vad.set(p.get("use_vad", False))
            self.vad_level.set(p.get("vad_level", "med"))
            self.word_ts.set(p.get("word_ts", False))
            self.deepl_key.set(p.get("deepl_key", ""))
            self.aai_key.set(p.get("aai_key", ""))
        except Exception as e:
            try:
                print(f"[PREF ERROR] Gagal load prefs: {e}")
            except Exception:
                pass

    # --- UI/UX BUILDER (V8.0) ---
    def setup_ui(self):
        t = self.i18n[self.current_lang]
        main_frm = ctk.CTkFrame(self)
        main_frm.pack(fill="both", expand=True, padx=16, pady=10)
        
        # TOP BAR (Enterprise Toolbar)
        top_bar = ctk.CTkFrame(main_frm, fg_color="transparent")
        top_bar.pack(fill="x", pady=(0, 4))
        ctk.CTkLabel(top_bar, text="Platinum Transcriber", font=("Arial", 18, "bold")).pack(side="left", padx=(4, 2))
        ctk.CTkLabel(top_bar, text="V8.0", font=("Arial", 11), text_color="#888888").pack(side="left", padx=(0, 8), pady=(4, 0))
        self.btn_lang = ctk.CTkButton(
            top_bar,
            text=(t["btn_lang_to_en"] if self.current_lang == "id" else t["btn_lang_to_id"]),
            command=self.toggle_lang,
            width=110, height=30, font=("Arial", 12),
        )
        self.btn_lang.pack(side="right", padx=3)
        self.btn_hist = ctk.CTkButton(top_bar, text=t["btn_hist"], command=self.show_history, width=110, height=30, font=("Arial", 12), fg_color="gray")
        self.btn_hist.pack(side="right", padx=3)
        self.btn_diag = ctk.CTkButton(top_bar, text=t["btn_diag"], fg_color="#ff9800", hover_color="#e68a00", command=self.run_diag, width=180, height=30, font=("Arial", 12))
        self.btn_diag.pack(side="right", padx=3)

        self.tabs = ctk.CTkTabview(main_frm); self.tabs.pack(fill="both", expand=True, pady=5)
        self.t1 = self.tabs.add(t["tab1"]); self.t2 = self.tabs.add(t["tab2"]); self.t3 = self.tabs.add(t["tab3"])

        # UI HELPER FUNCTIONS
        def _build_two_col(tab):
            root = ctk.CTkFrame(tab, fg_color="transparent")
            root.pack(fill="both", expand=True, padx=4, pady=4)
            root.grid_columnconfigure(0, weight=3, uniform="cols")
            root.grid_columnconfigure(1, weight=2, uniform="cols")
            left = ctk.CTkFrame(root, fg_color="transparent")
            right = ctk.CTkFrame(root, fg_color="transparent")
            left.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
            right.grid(row=0, column=1, sticky="nsew")
            left.grid_columnconfigure(0, weight=1)
            right.grid_columnconfigure(0, weight=1)
            return left, right

        def _section(parent, title):
            frm = ctk.CTkFrame(parent, corner_radius=12)
            frm.pack(fill="x", pady=(0, 10))
            frm.grid_columnconfigure(0, weight=1)
            ctk.CTkLabel(frm, text=title, font=("Arial", 14, "bold")).grid(row=0, column=0, sticky="w", padx=12, pady=(10, 6))
            body = ctk.CTkFrame(frm, fg_color="transparent")
            body.grid(row=1, column=0, sticky="nsew", padx=12, pady=(0, 10))
            body.grid_columnconfigure(0, weight=1)
            return body

        def _summary_panel(parent, title):
            card = ctk.CTkFrame(parent, corner_radius=12)
            card.pack(fill="both", expand=True)
            card.grid_columnconfigure(0, weight=1)
            ctk.CTkLabel(card, text=title, font=("Arial", 14, "bold"), text_color="#3B9ED9").grid(row=0, column=0, sticky="w", padx=12, pady=(10, 6))
            body = ctk.CTkFrame(card, fg_color="transparent")
            body.grid(row=1, column=0, sticky="nsew", padx=12, pady=(0, 12))
            body.grid_columnconfigure(0, weight=1)
            return card, body

        # ==========================================
        # TAB 1: INPUT & MESIN (Layout 2 Kolom)
        # ==========================================
        t1_left, t1_right = _build_two_col(self.t1)
        
        # Left Panel (Input)
        sec_src = _section(t1_left, t["sec_src"])
        ctk.CTkRadioButton(
            sec_src,
            text=t["src_yt"],
            variable=self.src_type,
            value="youtube",
            command=lambda: (self.toggle_src(), self.update_summary()),
        ).grid(row=0, column=0, sticky="w", pady=5)
        self.url_ent = ctk.CTkEntry(sec_src, textvariable=self.url_input, width=400, placeholder_text=t["ph_yt"])
        self.url_ent.grid(row=1, column=0, sticky="ew", pady=5)
        self.url_ent.bind("<KeyRelease>", self.update_summary_event)
        
        ctk.CTkRadioButton(
            sec_src,
            text=t["src_local"],
            variable=self.src_type,
            value="local",
            command=lambda: (self.toggle_src(), self.update_summary()),
        ).grid(row=2, column=0, sticky="w", pady=5)
        self.f_frm = ctk.CTkFrame(sec_src, fg_color="transparent")
        self.f_frm.grid(row=3, column=0, sticky="ew", pady=5)
        ctk.CTkButton(self.f_frm, text=t["browse"], command=self.browse).pack(side="left")
        self.f_lbl = ctk.CTkLabel(self.f_frm, text=t["no_file"], text_color="gray")
        self.f_lbl.pack(side="left", padx=10)

        sec_model = _section(t1_left, t["sec_model"])
        ctk.CTkRadioButton(sec_model, text=t["model_small"], variable=self.ai_eng, value="small", command=self.update_summary_event).grid(row=0, column=0, sticky="w", pady=5)
        ctk.CTkRadioButton(sec_model, text=t["model_medium"], variable=self.ai_eng, value="medium", command=self.update_summary_event).grid(row=1, column=0, sticky="w", pady=5)

        # Right Panel (Summary Tab 1)
        self.t1_summary_card, t1_sum_body = _summary_panel(t1_right, t["sum_exec"])
        self.lbl_sum_src = ctk.CTkLabel(t1_sum_body, text="", justify="left", font=("Arial", 13))
        self.lbl_sum_src.grid(row=0, column=0, sticky="w", pady=5)
        self.lbl_sum_model = ctk.CTkLabel(t1_sum_body, text="", justify="left", font=("Arial", 13))
        self.lbl_sum_model.grid(row=1, column=0, sticky="w", pady=5)

        # ==========================================
        # TAB 2: OUTPUT & TERJEMAHAN
        # ==========================================
        t2_left, t2_right = _build_two_col(self.t2)
        
        # Left Panel
        sec_out = _section(t2_left, t["sec_out"])
        frm_out = ctk.CTkFrame(sec_out, fg_color="transparent")
        frm_out.grid(row=0, column=0, sticky="w")
        row1_fmts = [(".srt", self.out_srt), (".vtt", self.out_vtt), (".txt", self.out_txt), (".docx", self.out_docx), (".pdf", self.out_pdf)]
        for text, var in row1_fmts:
            ctk.CTkCheckBox(frm_out, text=text, variable=var, command=self.update_summary_event).pack(side="left", padx=5, pady=5)

        frm_out2 = ctk.CTkFrame(sec_out, fg_color="transparent")
        frm_out2.grid(row=1, column=0, sticky="w")
        row2_fmts = [(".json", self.out_json), (".tsv", self.out_tsv), (".sbv", self.out_sbv), (".md", self.out_md), (".ass", self.out_ass)]
        for text, var in row2_fmts:
            ctk.CTkCheckBox(frm_out2, text=text, variable=var, command=self.update_summary_event).pack(side="left", padx=5, pady=5)

        frm_word_ts = ctk.CTkFrame(sec_out, fg_color="transparent")
        frm_word_ts.grid(row=2, column=0, sticky="w", pady=(2, 5))
        ctk.CTkCheckBox(frm_word_ts, text=t["use_word_ts"], variable=self.word_ts, command=self.update_summary_event).pack(side="left")
        ctk.CTkLabel(frm_word_ts, text=f"  ({t['word_ts_warn']})", font=("Arial", 10), text_color="#AAAAAA").pack(side="left")

        sec_trans = _section(t2_left, t["sec_trans"])
        ctk.CTkRadioButton(sec_trans, text=t["tr_none"], variable=self.trans_opt, value="none", command=self.update_translation_disclosure).grid(row=0, column=0, sticky="w", pady=5)
        ctk.CTkRadioButton(sec_trans, text=t["tr_google"], variable=self.trans_opt, value="google", command=self.update_translation_disclosure).grid(row=1, column=0, sticky="w", pady=5)
        ctk.CTkRadioButton(sec_trans, text=t["tr_any2en"], variable=self.trans_opt, value="any2en", command=self.update_translation_disclosure).grid(row=2, column=0, sticky="w", pady=5)
        ctk.CTkRadioButton(sec_trans, text=t["tr_deepl"], variable=self.trans_opt, value="deepl", command=self.update_translation_disclosure).grid(row=3, column=0, sticky="w", pady=5)
        
        self.deepl_row = ctk.CTkFrame(sec_trans, fg_color="transparent")
        self.deepl_row.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(self.deepl_row, text=t["lbl_deepl_key"], font=("Arial", 12, "bold"), text_color="#AAAAAA").grid(row=0, column=0, sticky="w", pady=(6, 0))
        self.deepl_ent = ctk.CTkEntry(self.deepl_row, textvariable=self.deepl_key, width=420, show="*", placeholder_text=t["ph_deepl"])
        self.deepl_ent.grid(row=1, column=0, sticky="ew", pady=(2, 5))

        # Right Panel (Summary Tab 2)
        self.t2_summary_card, t2_sum_body = _summary_panel(t2_right, t["sum_result"])
        self.lbl_sum_out = ctk.CTkLabel(t2_sum_body, text="", justify="left", font=("Arial", 13))
        self.lbl_sum_out.grid(row=0, column=0, sticky="w", pady=5)
        self.lbl_sum_tr = ctk.CTkLabel(t2_sum_body, text="", justify="left", font=("Arial", 13))
        self.lbl_sum_tr.grid(row=1, column=0, sticky="w", pady=5)

        # ==========================================
        # TAB 3: ASSEMBLY AI & CLOUD
        # ==========================================
        t3_left, t3_right = _build_two_col(self.t3)
        
        # Left Panel
        sec_pre = _section(t3_left, t["sec_pre"])
        ctk.CTkCheckBox(sec_pre, text=t["pre_ffmpeg"], variable=self.clean_audio, command=self.update_summary_event).grid(row=0, column=0, sticky="w", pady=5)
        ctk.CTkCheckBox(sec_pre, text=t["use_vad"], variable=self.use_vad, command=self._update_vad_disclosure).grid(row=1, column=0, sticky="w", pady=5)
        self._vad_level_row = ctk.CTkFrame(sec_pre, fg_color="transparent")
        self._vad_level_row.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(self._vad_level_row, text=t["vad_level"], font=("Arial", 12, "bold"), text_color="#AAAAAA").grid(row=0, column=0, sticky="w")
        self._vad_level_values = [t["vad_low"], t["vad_med"], t["vad_high"]]
        self._vad_level_menu = ctk.CTkOptionMenu(
            self._vad_level_row, values=self._vad_level_values,
            width=250, command=self._on_vad_level_change,
        )
        idx = {"low": 0, "med": 1, "high": 2}.get(self.vad_level.get(), 1)
        self._vad_level_menu.set(self._vad_level_values[idx])
        self._vad_level_menu.grid(row=1, column=0, sticky="w", pady=(2, 5))

        sec_cloud = _section(t3_left, t["sec_cloud"])
        ctk.CTkCheckBox(
            sec_cloud,
            text=t["use_aai"],
            variable=self.use_aai,
            command=self.update_aai_disclosure,
        ).grid(row=0, column=0, sticky="w", pady=5)
        
        self.aai_panel = ctk.CTkFrame(sec_cloud, corner_radius=10, fg_color="transparent")
        self.aai_panel.grid_columnconfigure(0, weight=1)
        frm_aai_opts = ctk.CTkFrame(self.aai_panel, fg_color="transparent")
        frm_aai_opts.grid(row=0, column=0, sticky="w", pady=5)
        ctk.CTkCheckBox(frm_aai_opts, text=t["aai_chap"], variable=self.aai_chap, command=self.update_summary_event).pack(side="left", padx=5)
        ctk.CTkCheckBox(frm_aai_opts, text=t["aai_pii"], variable=self.aai_pii, command=self.update_summary_event).pack(side="left", padx=5)
        ctk.CTkCheckBox(frm_aai_opts, text=t["aai_diar"], variable=self.aai_diar, command=self.update_summary_event).pack(side="left", padx=5)
        ctk.CTkCheckBox(frm_aai_opts, text=t["aai_filler"], variable=self.aai_filler, command=self.update_summary_event).pack(side="left", padx=5)
        ctk.CTkLabel(self.aai_panel, text=t["lbl_aai_key"], font=("Arial", 12, "bold"), text_color="#AAAAAA").grid(row=1, column=0, sticky="w", pady=(8, 0))
        ctk.CTkEntry(self.aai_panel, textvariable=self.aai_key, width=420, show="*", placeholder_text=t["ph_aai_key"]).grid(row=2, column=0, sticky="ew", pady=(2, 5))
        ctk.CTkLabel(self.aai_panel, text=t["lbl_vocab"], font=("Arial", 12, "bold"), text_color="#AAAAAA").grid(row=3, column=0, sticky="w", pady=(8, 0))
        ctk.CTkEntry(self.aai_panel, textvariable=self.aai_vocab, width=420, placeholder_text=t["ph_vocab_opt"]).grid(row=4, column=0, sticky="ew", pady=(2, 5))

        # Right Panel (Summary Tab 3)
        self.t3_summary_card, t3_sum_body = _summary_panel(t3_right, t["sum_cloud"])
        self.lbl_sum_cloud = ctk.CTkLabel(t3_sum_body, text="", justify="left", font=("Arial", 13))
        self.lbl_sum_cloud.grid(row=0, column=0, sticky="w", pady=5)

        # ==========================================
        # DASHBOARD BAWAH (Progress & Buttons)
        # ==========================================
        prog_frm = ctk.CTkFrame(main_frm, fg_color="transparent")
        prog_frm.pack(fill="x", pady=5)
        self.prog_bar = ctk.CTkProgressBar(prog_frm)
        self.prog_bar.pack(side="left", fill="x", expand=True, padx=(0, 10))
        self.prog_bar.set(0)
        self.prog_lbl = ctk.CTkLabel(prog_frm, text="0 %", font=("Arial", 12, "bold"), width=50)
        self.prog_lbl.pack(side="right")
        
        btn_frm = ctk.CTkFrame(main_frm, fg_color="transparent"); btn_frm.pack(fill="x", pady=5)
        btn_frm.grid_columnconfigure(0, weight=3)
        btn_frm.grid_columnconfigure(1, weight=1)
        btn_frm.grid_columnconfigure(2, weight=3)
        btn_frm.grid_columnconfigure(3, weight=1)
        
        self.btn_start = ctk.CTkButton(btn_frm, text=t["btn_start"], fg_color="#006400", hover_color="#004d00", font=("Arial", 15, "bold"), height=42, command=self.start_thread)
        self.btn_start.grid(row=0, column=0, sticky="ew", padx=(0, 6))

        self._live_eng_values = [t["live_engine_google"], t["live_engine_whisper"]]
        self.live_engine_menu = ctk.CTkOptionMenu(
            btn_frm, values=self._live_eng_values,
            width=160, height=42, font=("Arial", 11),
            command=self._on_live_engine_change,
        )
        self.live_engine_menu.set(self._live_eng_values[0])
        self.live_engine.set("google")
        self.live_engine_menu.grid(row=0, column=1, sticky="ew", padx=(0, 6))

        self.btn_live = ctk.CTkButton(
            btn_frm, text=t["btn_live"],
            fg_color="#7B2D8E", hover_color="#5E1F6E",
            font=("Arial", 15, "bold"), height=42,
            command=self._toggle_live_btn,
        )
        self.btn_live.grid(row=0, column=2, sticky="ew", padx=(0, 6))

        self.btn_stop = ctk.CTkButton(btn_frm, text=t["btn_stop"], fg_color="#8B0000", hover_color="#660000", font=("Arial", 13, "bold"), height=42, state="disabled", command=self.trigger_stop)
        self.btn_stop.grid(row=0, column=3, sticky="ew")
        
        self.console_log = ctk.CTkTextbox(main_frm, height=120, text_color="#D4D4D4", fg_color="#1E1E2E", font=("Consolas", 12))
        self.console_log.pack(fill="both", expand=True, pady=5); self.console_log.configure(state="disabled")
        
        # Keyboard shortcuts (Enterprise accessibility)
        self.bind("<Control-Return>", lambda ev: self.start_thread() if str(self.btn_start.cget("state")) != "disabled" else None)
        self.bind("<Escape>", lambda ev: self.trigger_stop() if str(self.btn_stop.cget("state")) != "disabled" else None)

        # Sinkronisasi awal progressive disclosure + ringkasan (UI only)
        self.toggle_src()
        self.update_translation_disclosure()
        self.update_aai_disclosure()
        self._update_vad_disclosure()
        self.update_summary_event()
        print("[SYSTEM] V8.0 Enterprise UI/UX Loaded.  |  Ctrl+Enter = Start  |  Esc = Stop")
        if HAS_FFMPEG: print("[SYSTEM] FFmpeg Portable Terdeteksi (Plug-and-Play Aktif).")

    # --- UI EVENT HANDLERS (DYNAMIC DISCLOSURE) ---
    def update_translation_disclosure(self, *args):
        if self.trans_opt.get() == "deepl": self.deepl_row.grid(row=4, column=0, sticky="ew", pady=5)
        else: self.deepl_row.grid_remove()
        self.update_summary()

    def _update_vad_disclosure(self, *args):
        if self.use_vad.get():
            self._vad_level_row.grid(row=2, column=0, sticky="ew", pady=5)
        else:
            self._vad_level_row.grid_remove()
        self.update_summary()

    def _on_vad_level_change(self, display_value):
        if display_value == self._vad_level_values[0]:
            self.vad_level.set("low")
        elif display_value == self._vad_level_values[2]:
            self.vad_level.set("high")
        else:
            self.vad_level.set("med")

    def update_aai_disclosure(self, *args):
        if self.use_aai.get(): self.aai_panel.grid(row=1, column=0, sticky="ew", pady=5)
        else: self.aai_panel.grid_remove()
        self.update_summary()

    def update_summary_event(self, *args):
        self.update_summary()

    def update_summary(self):
        """ Menulis struk ringkasan secara real-time di panel kanan """
        t = self.i18n[self.current_lang]
        src_mode = t["mode_yt"] if self.src_type.get() == "youtube" else t["mode_local"]
        src_val = self.url_input.get() if self.src_type.get() == "youtube" else t["target_files"].format(n=len(self.local_files))
        
        model_str = t["engine_small"] if self.ai_eng.get() == "small" else t["engine_medium"]
        if self.use_aai.get(): model_str = t["engine_cloud"]

        out_list = []
        for ext, var in [(".srt", self.out_srt), (".vtt", self.out_vtt), (".txt", self.out_txt),
                         (".docx", self.out_docx), (".pdf", self.out_pdf), (".json", self.out_json),
                         (".tsv", self.out_tsv), (".sbv", self.out_sbv), (".md", self.out_md), (".ass", self.out_ass)]:
            if var.get():
                out_list.append(ext)
        out_str = ", ".join(out_list) if out_list else t["out_none"]

        tr_dict = {"none": t["tr_map_none"], "google": t["tr_map_google"], "any2en": t["tr_map_any2en"], "deepl": t["tr_map_deepl"]}
        tr_str = tr_dict.get(self.trans_opt.get(), "")

        # Pembaruan Teks Dinamis
        self.lbl_sum_src.configure(text=t["sum_mode"].format(mode=src_mode, target=src_val))
        self.lbl_sum_model.configure(text=t["sum_engine"].format(engine=model_str))
        self.lbl_sum_out.configure(text=t["sum_out"].format(out=out_str))
        self.lbl_sum_tr.configure(text=t["sum_tr"].format(tr=tr_str))
        
        cloud_info = t["cloud_off"]
        if self.use_aai.get():
            opts = []
            if self.aai_chap.get(): opts.append(t["opt_chap"])
            if self.aai_pii.get(): opts.append(t["opt_pii"])
            if self.aai_diar.get(): opts.append(t["opt_diar"])
            cloud_info = t["cloud_on"].format(opts=(", ".join(opts) if opts else t["cloud_std"]))
        self.lbl_sum_cloud.configure(text=cloud_info)

    def toggle_lang(self):
        # Rebuild UI agar semua teks (termasuk judul tab) ikut berubah konsisten
        self.current_lang = "en" if self.current_lang == "id" else "id"
        for child in list(self.winfo_children()):
            try:
                child.destroy()
            except Exception:
                pass
        self.setup_ui()
        self.update_translation_disclosure()
        self.update_aai_disclosure()
        self._update_vad_disclosure()
        self.update_summary()

    def show_history(self):
        hist_win = ctk.CTkToplevel(self)
        t = self.i18n[self.current_lang]
        hist_win.title(t["hist_title"])
        hist_win.geometry("600x400")
        hist_win.attributes("-topmost", True)
        lbl = ctk.CTkLabel(hist_win, text=t["hist_lbl"], font=("Arial", 16, "bold"))
        lbl.pack(pady=10)
        textbox = ctk.CTkTextbox(hist_win, width=550, height=300)
        textbox.pack(padx=10, pady=10)
        try:
            conn = sqlite3.connect(DB_FILE)
            try:
                cursor = conn.execute("SELECT tanggal, file_name, durasi_detik, engine FROM logs ORDER BY id DESC")
                for row in cursor:
                    mins, secs = divmod(int(row[2]), 60)
                    textbox.insert("end", f"[{row[0]}] {row[1]}\nEngine: {row[3]} | Waktu Proses: {mins}m {secs}s\n{'-'*50}\n")
            finally:
                conn.close()
        except Exception:
            textbox.insert("end", "Belum ada riwayat / Database Error.")
        textbox.configure(state="disabled")

    def toggle_src(self):
        if self.src_type.get() == "youtube": self.f_frm.grid_remove(); self.url_ent.grid(row=1, column=0, sticky="ew", pady=5)
        else: self.url_ent.grid_remove(); self.f_frm.grid(row=3, column=0, sticky="ew", pady=5)

    _AUDIO_VIDEO_TYPES = [
        ("Audio/Video",
         "*.mp3 *.wav *.m4a *.ogg *.flac *.aac *.wma *.opus *.amr *.aiff *.aif *.mka *.caf *.ac3 "
         "*.mp4 *.mkv *.mov *.avi *.webm *.3gp *.3g2 *.ts *.mts *.m2ts *.mpg *.mpeg *.wmv *.flv *.vob"),
        ("Audio", "*.mp3 *.wav *.m4a *.ogg *.flac *.aac *.wma *.opus *.amr *.aiff *.aif *.mka *.caf *.ac3"),
        ("Video", "*.mp4 *.mkv *.mov *.avi *.webm *.3gp *.3g2 *.ts *.mts *.m2ts *.mpg *.mpeg *.wmv *.flv *.vob"),
        ("All files", "*.*"),
    ]

    def browse(self):
        f = filedialog.askopenfilenames(filetypes=self._AUDIO_VIDEO_TYPES)
        if f:
            self.local_files = list(f)
            t = self.i18n[self.current_lang]
            self.f_lbl.configure(text=t["file_count"].format(n=len(f)), text_color="#43A047")
            self.update_summary()

    def run_diag(self):
        ram = psutil.virtual_memory()
        fr = ram.available / (1024**3)
        tr = ram.total / (1024**3)
        fs = shutil.disk_usage(os.path.expanduser("~"))[2] / (1024**3)
        if HAS_FFMPEG:
            ff = "✅ (Portable)"
        else:
            ff = "✅ (System)" if _ensure_ffmpeg_available() else "❌"
        try:
            urllib.request.urlopen('http://google.com', timeout=3)
            inet = "✅"
        except Exception:
            inet = "❌"
        t = self.i18n[self.current_lang]
        messagebox.showinfo(t["mb_title_diag"], t["mb_msg_diag"].format(inet=inet, ff=ff, fr=fr, tr=tr, fs=fs))

    def format_ts(self, seconds, fmt="txt"):
        hours, remainder = divmod(seconds, 3600)
        minutes, secs = divmod(remainder, 60)
        if fmt in ["srt", "vtt"]:
            delim = "," if fmt == "srt" else "."
            ms = int((seconds - int(seconds)) * 1000)
            return f"{int(hours):02d}:{int(minutes):02d}:{int(secs):02d}{delim}{ms:03d}"
        else:
            if hours > 0: return f"[{int(hours):02d}:{int(minutes):02d}:{int(secs):02d}]"
            return f"[{int(minutes):02d}:{int(secs):02d}]"

    def yt_hook(self, d):
        if self.stop_flag:
            raise Exception("Interrupted by E-Stop!")
        if d['status'] == 'downloading':
            p = d.get('_percent_str', '0%').replace('%', '')
            try:
                self.msg_queue.put({"type": "progress", "val": float(p) / 100})
            except (ValueError, TypeError):
                pass

    def trigger_stop(self):
        self.stop_flag = True; print("\n[E-STOP] Menerima sinyal pembatalan. Menutup operasi dengan aman...")
        t = self.i18n[self.current_lang]
        self.btn_stop.configure(state="disabled", text=t["btn_stop_busy"])

    # --- LIVE DICTATION ENGINE ---
    def _toggle_live_btn(self):
        if self.live_recording:
            self.stop_live()
        else:
            self.start_live()

    def _on_live_engine_change(self, display_value):
        if display_value == self._live_eng_values[0]:
            self.live_engine.set("google")
        else:
            self.live_engine.set("whisper")

    def start_live(self):
        t = self.i18n[self.current_lang]
        if not HAS_PYAUDIO:
            messagebox.showerror(t["mb_title_error"], t["live_no_pyaudio"])
            return
        if not HAS_NUMPY:
            messagebox.showerror(t["mb_title_error"], t["live_no_numpy"])
            return

        engine = self.live_engine.get()
        if engine == "google" and not HAS_SPEECH_REC:
            messagebox.showerror(t["mb_title_error"], t["live_no_speech"])
            return

        self.live_recording = True
        self.live_stop_event.clear()
        self.live_transcript.clear()
        while not self.live_audio_queue.empty():
            try: self.live_audio_queue.get_nowait()
            except queue.Empty: break

        self.btn_start.configure(state="disabled")
        self.btn_live.configure(text=t["btn_live_stop"], fg_color="#8B0000", hover_color="#660000")
        self.live_engine_menu.configure(state="disabled")
        self.btn_lang.configure(state="disabled")
        self.btn_stop.configure(state="disabled")
        self.msg_queue.put({"type": "prog_mode", "val": "indeterminate"})

        print(t["live_starting"])

        raw_mode = (engine == "google")
        threading.Thread(target=self._live_producer_thread, args=(raw_mode,), daemon=True).start()
        if engine == "whisper":
            threading.Thread(target=self._live_consumer_whisper, daemon=True).start()
        else:
            threading.Thread(target=self._live_consumer_google, daemon=True).start()

    def stop_live(self):
        t = self.i18n[self.current_lang]
        self.live_stop_event.set()
        self._reset_live_ui_state()

        print(t["live_stopped"])

        with self._live_lock:
            transcript_copy = list(self.live_transcript)

        if not transcript_copy:
            print(t["live_empty"])
            return

        try:
            desktop = get_windows_desktop_dir()
            root_out = os.path.join(desktop, "PlatinumTranscriber_Output")
            os.makedirs(root_out, exist_ok=True)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            fname = f"LiveDictation_{ts}.txt"
            fpath = os.path.join(root_out, fname)

            engine_label = "Google Speech" if self.live_engine.get() == "google" else "Whisper (tiny)"
            full_text = " ".join(transcript_copy)

            with open(fpath, "w", encoding="utf-8") as f:
                f.write("=" * 60 + "\n")
                f.write("  LIVE DICTATION TRANSCRIPT\n")
                f.write(f"  Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
                f.write(f"  Engine: {engine_label}\n")
                f.write("=" * 60 + "\n\n")
                f.write(full_text + "\n\n")
                f.write("=" * 60 + "\n")

            print(t["live_saved"].format(path=fpath))
            save_log("LiveDictation", 0, engine_label)
            self.ui_call(messagebox.showinfo, t["mb_title_success"], t["live_saved"].format(path=fpath))
            self.ui_call(self.open_folder, root_out)
        except Exception as e:
            print(f"[LIVE SAVE ERROR] {e}")

    def _live_producer_thread(self, raw_mode=False):
        p = None
        stream = None
        try:
            p = pyaudio.PyAudio()
            try:
                dev = p.get_default_input_device_info()
                print(f"[MIC] Perangkat: {dev['name']} | {LIVE_RATE}Hz | Blok {LIVE_RECORD_SECONDS}s")
            except Exception:
                print("[MIC] Menggunakan perangkat input default")

            stream = p.open(
                format=pyaudio.paInt16, channels=1, rate=LIVE_RATE,
                input=True, frames_per_buffer=LIVE_CHUNK,
            )

            t = self.i18n[self.current_lang]
            print(t["live_mic_ready"])

            chunks_per_block = int(LIVE_RATE / LIVE_CHUNK * LIVE_RECORD_SECONDS)

            while not self.live_stop_event.is_set():
                frames = []
                for _ in range(chunks_per_block):
                    if self.live_stop_event.is_set():
                        break
                    try:
                        data = stream.read(LIVE_CHUNK, exception_on_overflow=False)
                        frames.append(data)
                    except Exception:
                        break

                if not frames:
                    continue

                raw_bytes = b"".join(frames)
                audio_f32 = np_audio.frombuffer(raw_bytes, dtype=np_audio.int16).astype(np_audio.float32) / 32768.0
                energy = _rms_energy(audio_f32)

                if energy < LIVE_SILENCE_THRESH:
                    continue

                if raw_mode:
                    self.live_audio_queue.put(raw_bytes)
                else:
                    self.live_audio_queue.put(audio_f32)
        except Exception as e:
            print(f"[MIC ERROR] {e}")
            self.live_stop_event.set()
            self.ui_call(self._reset_live_ui_state)
        finally:
            if stream is not None:
                try:
                    stream.stop_stream()
                    stream.close()
                except Exception:
                    pass
            if p is not None:
                try:
                    p.terminate()
                except Exception:
                    pass

    def _live_consumer_whisper(self):
        try:
            t = self.i18n[self.current_lang]
            print(t["live_warmup"])
            model = WhisperModel("tiny", device="cpu", compute_type="int8")

            dummy = np_audio.zeros(LIVE_RATE, dtype=np_audio.float32)
            list(model.transcribe(dummy, beam_size=1, language="id")[0])
            print(t["live_mic_ready"])

            while not self.live_stop_event.is_set():
                try:
                    audio_f32 = self.live_audio_queue.get(timeout=0.5)
                except queue.Empty:
                    continue

                while not self.live_audio_queue.empty():
                    try:
                        audio_f32 = self.live_audio_queue.get_nowait()
                    except queue.Empty:
                        break

                segments, _ = model.transcribe(audio_f32, beam_size=1, language="id")
                text = "".join([s.text for s in segments]).strip()

                if text:
                    with self._live_lock:
                        self.live_transcript.append(text)
                    self.msg_queue.put({"type": "live_text", "text": f"  >> {text}"})
        except Exception as e:
            print(f"[WHISPER LIVE ERROR] {e}")
            self.live_stop_event.set()
            self.ui_call(self._reset_live_ui_state)

    def _live_consumer_google(self):
        try:
            recognizer = sr.Recognizer()
            t = self.i18n[self.current_lang]
            print(t["live_mic_ready"])

            while not self.live_stop_event.is_set():
                try:
                    raw_pcm = self.live_audio_queue.get(timeout=0.5)
                except queue.Empty:
                    continue

                while not self.live_audio_queue.empty():
                    try:
                        raw_pcm = self.live_audio_queue.get_nowait()
                    except queue.Empty:
                        break

                try:
                    audio_data = sr.AudioData(raw_pcm, LIVE_RATE, 2)
                    text = recognizer.recognize_google(audio_data, language="id-ID")
                    if text and text.strip():
                        with self._live_lock:
                            self.live_transcript.append(text.strip())
                        self.msg_queue.put({"type": "live_text", "text": f"  >> {text.strip()}"})
                except sr.UnknownValueError:
                    pass
                except sr.RequestError as e:
                    print(f"[GOOGLE ERROR] {e}")
                except Exception as e:
                    print(f"[LIVE ERROR] {e}")
        except Exception as e:
            print(f"[GOOGLE LIVE ERROR] {e}")
            self.live_stop_event.set()
            self.ui_call(self._reset_live_ui_state)

    # --- LOGIKA MESIN (BACKEND TIDAK DIUBAH, 100% AMAN) ---
    def start_thread(self):
        # Snapshot state UI di main thread (hindari akses Tk Variable dari thread belakang)
        self.save_prefs()
        self.stop_flag = False
        snapshot = {
            "src_type": self.src_type.get(),
            "url_input": self.url_input.get(),
            "local_files": list(self.local_files) if self.local_files else [],
            "ai_eng": self.ai_eng.get(),
            "clean_audio": bool(self.clean_audio.get()),
            "trans_opt": self.trans_opt.get(),
            "deepl_key": self.deepl_key.get(),
            "out_srt": bool(self.out_srt.get()),
            "out_vtt": bool(self.out_vtt.get()),
            "out_txt": bool(self.out_txt.get()),
            "out_docx": bool(self.out_docx.get()),
            "out_pdf": bool(self.out_pdf.get()),
            "out_json": bool(self.out_json.get()),
            "out_tsv": bool(self.out_tsv.get()),
            "out_sbv": bool(self.out_sbv.get()),
            "out_md": bool(self.out_md.get()),
            "out_ass": bool(self.out_ass.get()),
            "use_aai": bool(self.use_aai.get()),
            "aai_key": self.aai_key.get(),
            "aai_vocab": self.aai_vocab.get(),
            "aai_chap": bool(self.aai_chap.get()),
            "aai_pii": bool(self.aai_pii.get()),
            "aai_diar": bool(self.aai_diar.get()),
            "aai_filler": bool(self.aai_filler.get()),
            "use_vad": bool(self.use_vad.get()),
            "vad_level": self.vad_level.get(),
            "word_ts": bool(self.word_ts.get()),
        }
        threading.Thread(target=self.logic, args=(snapshot,), daemon=True).start()

    def logic(self, snapshot):
        self.ui_call(self.btn_start.configure, state="disabled")
        self.ui_call(self.btn_live.configure, state="disabled")
        self.ui_call(self.live_engine_menu.configure, state="disabled")
        self.ui_call(self.btn_lang.configure, state="disabled")
        self.ui_call(self.btn_stop.configure, state="normal", text=self.i18n[self.current_lang]["btn_stop"])
        self.msg_queue.put({"type": "progress", "val": 0})
        self.msg_queue.put({"type": "prog_mode", "val": "determinate"})
        _result = ("error", "Unexpected termination")

        try:
            ram_free = psutil.virtual_memory().available / (1024**3)
            desktop = get_windows_desktop_dir()
            root_out = os.path.join(desktop, "PlatinumTranscriber_Output")
            run_folder = os.path.join(root_out, datetime.datetime.now().strftime('%Y%m%d_%H%M%S'))
            os.makedirs(run_folder, exist_ok=True)
            files_to_do = []
            temp_mp3: Optional[str] = None

            if snapshot["src_type"] == "youtube":
                if not snapshot["url_input"]:
                    raise ValueError("URL YouTube Kosong!")
                print("[YT] Mengekstrak audio jalur aman...")
                temp_mp3 = os.path.join(run_folder, "temp.mp3")
                with yt_dlp.YoutubeDL({'format': 'bestaudio', 'outtmpl': temp_mp3, 'progress_hooks': [self.yt_hook], 'quiet': True}) as ydl:
                    ydl.download([snapshot["url_input"]])
                files_to_do.append((temp_mp3, "YouTube_Extract"))
            else:
                if not snapshot["local_files"]:
                    raise ValueError("Belum ada file lokal yang dipilih!")
                for f in snapshot["local_files"]:
                    files_to_do.append((f, os.path.splitext(os.path.basename(f))[0]))

            t_loc = self.i18n[self.current_lang]
            total_bytes = 0
            ext_set = set()
            for path_check, _ in files_to_do:
                try:
                    fsize = os.path.getsize(path_check)
                    total_bytes += fsize
                    ext_set.add(os.path.splitext(path_check)[1].lower() or "?")
                    if fsize > LARGE_FILE_WARN_GB * (1024 ** 3):
                        print(t_loc["warn_large_file"].format(
                            name=os.path.basename(path_check),
                            size=fsize / (1024 ** 3),
                        ))
                except OSError:
                    pass

            if total_bytes >= 1024 ** 3:
                total_fmt = f"{total_bytes / (1024 ** 3):.2f} GB"
            elif total_bytes >= 1024 ** 2:
                total_fmt = f"{total_bytes / (1024 ** 2):.1f} MB"
            else:
                total_fmt = f"{total_bytes / 1024:.0f} KB"
            print(t_loc["fmt_total"].format(
                n=len(files_to_do),
                total_fmt=total_fmt,
                exts=", ".join(sorted(ext_set)) if ext_set else "-",
            ))

            engine = snapshot["trans_opt"]
            deepl_tr, google_tr, any_tr = None, None, None
            if engine == "deepl":
                key = (snapshot["deepl_key"] or "").strip()
                if not key:
                    print("[API WARNING] DeepL API Key kosong. Fallback ke Google Translate.")
                    engine = "google"
                else:
                    try:
                        deepl_tr = deepl.Translator(key)
                    except Exception as e:
                        engine = "google"
                        print(f"[API ERROR] DeepL Error ({e}). Fallback ke Google Translate.")
            
            if engine == "google": google_tr = GoogleTranslator(source='auto', target='id')
            elif engine == "any2en": any_tr = GoogleTranslator(source='auto', target='en')

            if snapshot["use_aai"]:
                aai_key = (snapshot["aai_key"] or "").strip()
                if not aai_key:
                    raise ValueError("AssemblyAI API Key kosong.")
                
                print("[SISTEM] Menghubungkan ke AssemblyAI Cloud...")
                aai.settings.api_key = aai_key
                vocab = [w.strip() for w in (snapshot["aai_vocab"] or "").split(",") if w.strip()] or None
                config = aai.TranscriptionConfig(
                    speech_models=["universal-3-pro"],
                    auto_chapters=snapshot["aai_chap"],
                    entity_detection=snapshot["aai_pii"],
                    speaker_labels=snapshot["aai_diar"],
                    disfluencies=not snapshot["aai_filler"],
                    word_boost=vocab,
                )
                transcriber = aai.Transcriber(config=config)
                self.msg_queue.put({"type": "prog_mode", "val": "indeterminate"})
                
                for path, name in files_to_do:
                    if self.stop_flag: break
                    print(f"[CLOUD] Memproses dokumen: {name}...")
                    t_start = time.time()
                    transcript = transcriber.transcribe(path)
                    if transcript.error: raise ValueError(transcript.error)

                    bp = os.path.join(run_folder, name)

                    aai_engine = "AssemblyAI Cloud (universal-3-pro)"
                    aai_paragraphs = _text_to_paragraphs(transcript.text)

                    if snapshot["out_srt"]:
                        try:
                            srt_content = transcript.export_subtitles_srt()
                            with open(f"{bp}.srt", "w", encoding="utf-8") as f_out:
                                f_out.write(srt_content)
                            print("[RENDER] File .srt sukses dibuat.")
                        except Exception as srt_exc:
                            print(f"[SRT WARNING] Gagal export SRT: {srt_exc}")

                    if snapshot["out_vtt"]:
                        try:
                            vtt_content = transcript.export_subtitles_vtt()
                            with open(f"{bp}.vtt", "w", encoding="utf-8") as f_out:
                                f_out.write(vtt_content)
                            print("[RENDER] File .vtt sukses dibuat.")
                        except Exception as vtt_exc:
                            print(f"[VTT WARNING] Gagal export VTT: {vtt_exc}")

                    if snapshot["out_txt"]:
                        with open(f"{bp}.txt", "w", encoding="utf-8") as f_out:
                            f_out.write(f"{'═' * 60}\n")
                            f_out.write(f"  TRANSCRIPTION: {name}\n")
                            f_out.write(f"  Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
                            f_out.write(f"  Engine: {aai_engine}\n")
                            f_out.write(f"{'═' * 60}\n\n")
                            for para in aai_paragraphs:
                                f_out.write(f"{para}\n\n")
                        print("[RENDER] File .txt sukses dibuat.")

                    if snapshot["out_docx"]:
                        try:
                            _render_docx(f"{bp}.docx", name, aai_engine, aai_paragraphs)
                            print("[RENDER] File .docx sukses dibuat.")
                        except Exception as dx_exc:
                            print(f"[Word Error] {dx_exc}")

                    if snapshot["out_pdf"]:
                        try:
                            _render_pdf(f"{bp}.pdf", name, aai_engine, aai_paragraphs)
                            print("[RENDER] File .pdf sukses dibuat.")
                        except Exception as pdf_exc:
                            print(f"[PDF Error] {pdf_exc}")

                    if snapshot["out_md"]:
                        with open(f"{bp}.md", "w", encoding="utf-8") as f_out:
                            f_out.write(f"# {name}\n\n")
                            f_out.write(f"**Date:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  \n")
                            f_out.write(f"**Engine:** {aai_engine}\n\n---\n\n")
                            for para in aai_paragraphs:
                                f_out.write(f"{para}\n\n")
                        print("[RENDER] File .md sukses dibuat.")

                    aai_words = getattr(transcript, "words", None) or []
                    aai_utterances = getattr(transcript, "utterances", None) or []
                    aai_subs = aai_utterances if aai_utterances else [type("U", (), {"start": w.start, "end": w.end, "text": w.text})() for w in aai_words]

                    if snapshot["out_json"]:
                        jdata = {"file": name, "engine": aai_engine,
                                 "date": datetime.datetime.now().isoformat(),
                                 "segments": [{"start": round(u.start / 1000, 3), "end": round(u.end / 1000, 3), "text": u.text} for u in aai_subs]}
                        with open(f"{bp}.json", "w", encoding="utf-8") as f_out:
                            json.dump(jdata, f_out, ensure_ascii=False, indent=2)
                        print("[RENDER] File .json sukses dibuat.")

                    if snapshot["out_tsv"]:
                        with open(f"{bp}.tsv", "w", encoding="utf-8") as f_out:
                            f_out.write("start\tend\ttext\n")
                            for u in aai_subs:
                                f_out.write(f"{u.start / 1000:.3f}\t{u.end / 1000:.3f}\t{u.text}\n")
                        print("[RENDER] File .tsv sukses dibuat.")

                    if snapshot["out_sbv"]:
                        with open(f"{bp}.sbv", "w", encoding="utf-8") as f_out:
                            for u in aai_subs:
                                f_out.write(f"{_fmt_sbv_ts(u.start / 1000)},{_fmt_sbv_ts(u.end / 1000)}\n{u.text}\n\n")
                        print("[RENDER] File .sbv sukses dibuat.")

                    if snapshot["out_ass"]:
                        with open(f"{bp}.ass", "w", encoding="utf-8") as f_out:
                            f_out.write(_ass_header())
                            for u in aai_subs:
                                f_out.write(f"Dialogue: 0,{_fmt_ass_ts(u.start / 1000)},{_fmt_ass_ts(u.end / 1000)},Default,,0,0,0,,{u.text}\n")
                        print("[RENDER] File .ass sukses dibuat.")

                    if snapshot["aai_chap"] and transcript.chapters:
                        with open(os.path.join(run_folder, f"{name}_Chapters.txt"), "w", encoding="utf-8") as f_chap:
                            f_chap.write(f"{'═' * 60}\n")
                            f_chap.write(f"  AUTO CHAPTERS: {name}\n")
                            f_chap.write(f"  Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
                            f_chap.write(f"  Engine: {aai_engine}\n")
                            f_chap.write(f"{'═' * 60}\n\n")
                            for idx, ch in enumerate(transcript.chapters, start=1):
                                f_chap.write(f"── Chapter {idx} [{_fmt_ms(ch.start)} - {_fmt_ms(ch.end)}] ──\n")
                                f_chap.write(f"  {ch.headline}\n\n")
                                f_chap.write(f"  {ch.summary}\n\n")
                        print("[RENDER] File Chapters sukses dibuat.")

                    save_log(name, time.time() - t_start, "AssemblyAI Cloud")
                self.msg_queue.put({"type": "prog_mode", "val": "determinate"})

            else:
                m_type = snapshot["ai_eng"]
                if ram_free < MIN_FREE_RAM_GB:
                    raise ValueError(f"RAM Anda terlalu kecil. Minimal {MIN_FREE_RAM_GB}GB Free RAM.")
                print(f"[AI ENGINE] Memuat arsitektur saraf FASTER-WHISPER {m_type.upper()}...")
                self.msg_queue.put({"type": "prog_mode", "val": "indeterminate"})
                
                model = WhisperModel(m_type, device="cpu", compute_type="int8")
                
                for path, name in files_to_do:
                    if self.stop_flag: break
                    t_start = time.time()
                    proc_path = path
                    if snapshot["clean_audio"]:
                        if not _ensure_ffmpeg_available():
                            raise ValueError("FFmpeg tidak ditemukan. Matikan 'Noise Reduction' atau pasang FFmpeg/sertakan ffmpeg.exe di folder aplikasi.")
                        clean_path = os.path.join(run_folder, f"clean_{os.path.basename(path)}.wav")
                        r = subprocess.run(
                            ["ffmpeg", "-y", "-i", path, "-af", "afftdn", clean_path],
                            stdout=subprocess.DEVNULL,
                            stderr=subprocess.DEVNULL,
                        )
                        if r.returncode != 0 or (not os.path.exists(clean_path)):
                            raise ValueError("FFmpeg gagal melakukan noise reduction. Cek input audio atau instalasi FFmpeg.")
                        proc_path = clean_path

                    print(f"\n[AI] Menganalisis gelombang suara: {name}")
                    self.msg_queue.put({"type": "prog_mode", "val": "determinate"})

                    transcribe_kwargs = {"beam_size": 5}
                    if snapshot["use_vad"]:
                        vp = VAD_PRESETS.get(snapshot["vad_level"], VAD_PRESETS["med"])
                        transcribe_kwargs["vad_filter"] = True
                        transcribe_kwargs["vad_parameters"] = vp
                        print(f"[VAD] Filter aktif (level={snapshot['vad_level']}, threshold={vp['threshold']})")
                    use_word = snapshot["word_ts"] and snapshot["trans_opt"] == "none"
                    if use_word:
                        transcribe_kwargs["word_timestamps"] = True
                        print("[WORD-TS] Word-level timestamps aktif")

                    segments, info = model.transcribe(proc_path, **transcribe_kwargs)
                    audio_duration = info.duration
                    
                    bp = os.path.join(run_folder, name)
                    doc_lines = []
                    
                    whisper_engine = f"Faster-Whisper ({m_type.capitalize()})"
                    print(f"[I/O STREAM] Membuka jalur penulisan seketika ke Hardisk...")
                    f_srt = open(f"{bp}.srt", "w", encoding="utf-8") if snapshot["out_srt"] else None
                    f_vtt = open(f"{bp}.vtt", "w", encoding="utf-8") if snapshot["out_vtt"] else None
                    f_txt = open(f"{bp}.txt", "w", encoding="utf-8") if snapshot["out_txt"] else None
                    f_sbv = open(f"{bp}.sbv", "w", encoding="utf-8") if snapshot["out_sbv"] else None
                    f_tsv = open(f"{bp}.tsv", "w", encoding="utf-8") if snapshot["out_tsv"] else None
                    f_md  = open(f"{bp}.md",  "w", encoding="utf-8") if snapshot["out_md"]  else None
                    f_ass = open(f"{bp}.ass", "w", encoding="utf-8") if snapshot["out_ass"] else None
                    json_segments = [] if snapshot["out_json"] else None
                    if f_vtt: f_vtt.write("WEBVTT\n\n")
                    if f_tsv: f_tsv.write("start\tend\ttext\n")
                    if f_ass: f_ass.write(_ass_header())
                    if f_md:
                        f_md.write(f"# {name}\n\n")
                        f_md.write(f"**Date:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  \n")
                        f_md.write(f"**Engine:** {whisper_engine}\n\n---\n\n")
                    if f_txt:
                        f_txt.write(f"{'═' * 60}\n")
                        f_txt.write(f"  TRANSCRIPTION: {name}\n")
                        f_txt.write(f"  Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
                        f_txt.write(f"  Engine: {whisper_engine}\n")
                        f_txt.write(f"{'═' * 60}\n\n")

                    try:
                        sub_idx = 0
                        for i, seg in enumerate(segments, start=1):
                            if self.stop_flag: break
                            if audio_duration > 0: self.msg_queue.put({"type": "progress", "val": seg.end / audio_duration})

                            s, e, text = seg.start, seg.end, seg.text.strip()
                            try:
                                if engine == "google": text = google_tr.translate(text)
                                elif engine == "deepl": text = deepl_tr.translate_text(text, target_lang="ID").text
                                elif engine == "any2en": text = any_tr.translate(text)
                            except Exception as tr_exc:
                                text = f"[Translate Error] {text}"
                                print(f"[TRANSLATE ERROR] {tr_exc}")

                            if use_word and getattr(seg, "words", None):
                                for w in seg.words:
                                    wt = w.word.strip()
                                    if not wt:
                                        continue
                                    sub_idx += 1
                                    if f_srt: f_srt.write(f"{sub_idx}\n{self.format_ts(w.start, 'srt')} --> {self.format_ts(w.end, 'srt')}\n{wt}\n\n")
                                    if f_vtt: f_vtt.write(f"{self.format_ts(w.start, 'vtt')} --> {self.format_ts(w.end, 'vtt')}\n{wt}\n\n")
                                    if f_sbv: f_sbv.write(f"{_fmt_sbv_ts(w.start)},{_fmt_sbv_ts(w.end)}\n{wt}\n\n")
                                    if f_ass: f_ass.write(f"Dialogue: 0,{_fmt_ass_ts(w.start)},{_fmt_ass_ts(w.end)},Default,,0,0,0,,{wt}\n")
                                    if f_tsv: f_tsv.write(f"{w.start:.3f}\t{w.end:.3f}\t{wt}\n")
                                    if json_segments is not None:
                                        json_segments.append({"start": round(w.start, 3), "end": round(w.end, 3), "text": wt})
                            else:
                                sub_idx += 1
                                if f_srt: f_srt.write(f"{sub_idx}\n{self.format_ts(s, 'srt')} --> {self.format_ts(e, 'srt')}\n{text}\n\n")
                                if f_vtt: f_vtt.write(f"{self.format_ts(s, 'vtt')} --> {self.format_ts(e, 'vtt')}\n{text}\n\n")
                                if f_sbv: f_sbv.write(f"{_fmt_sbv_ts(s)},{_fmt_sbv_ts(e)}\n{text}\n\n")
                                if f_ass: f_ass.write(f"Dialogue: 0,{_fmt_ass_ts(s)},{_fmt_ass_ts(e)},Default,,0,0,0,,{text}\n")
                                if f_tsv: f_tsv.write(f"{s:.3f}\t{e:.3f}\t{text}\n")
                                if json_segments is not None:
                                    json_segments.append({"start": round(s, 3), "end": round(e, 3), "text": text})

                            txt_line = f"{self.format_ts(s, 'txt')} {text}"
                            if f_txt: f_txt.write(f"{txt_line}\n")
                            if f_md: f_md.write(f"**[{self.format_ts(s, 'txt')}]** {text}  \n")
                            if snapshot["out_docx"] or snapshot["out_pdf"]:
                                doc_lines.append(txt_line)
                            if i % 10 == 0: print(f"Menulis baris ke-{i}...")
                    finally:
                        for fh in (f_srt, f_vtt, f_txt, f_sbv, f_tsv, f_md, f_ass):
                            if fh:
                                fh.close()

                    if self.stop_flag: break

                    if json_segments is not None:
                        with open(f"{bp}.json", "w", encoding="utf-8") as f_json:
                            json.dump({"file": name, "engine": whisper_engine,
                                       "date": datetime.datetime.now().isoformat(),
                                       "segments": json_segments}, f_json, ensure_ascii=False, indent=2)
                        print("[RENDER] File .json sukses dibuat.")

                    if snapshot["out_docx"]:
                        try:
                            _render_docx(f"{bp}.docx", name, whisper_engine, doc_lines)
                            print("[RENDER] File .docx sukses dibuat.")
                        except Exception as dx_exc:
                            print(f"[Word Error] {dx_exc}")

                    if snapshot["out_pdf"]:
                        try:
                            _render_pdf(f"{bp}.pdf", name, whisper_engine, doc_lines)
                            print("[RENDER] File .pdf sukses dibuat.")
                        except Exception as pdf_exc:
                            print(f"[PDF Error] {pdf_exc}")
                    
                    if snapshot["clean_audio"] and os.path.exists(proc_path):
                        try:
                            os.remove(proc_path)
                        except Exception as e:
                            print(f"[CLEANUP WARNING] Gagal hapus file sementara: {e}")
                    save_log(name, time.time() - t_start, f"Faster-Whisper ({m_type.capitalize()})")

            if snapshot["src_type"] == "youtube":
                try:
                    if os.path.exists(temp_mp3):
                        os.remove(temp_mp3)
                except Exception as e:
                    print(f"[CLEANUP WARNING] Gagal hapus temp.mp3: {e}")

            if self.stop_flag:
                _result = ("cancelled", None)
            else:
                _result = ("success", run_folder)

        except Exception as e:
            _result = ("error", str(e))
        finally:
            self.ui_call(self.btn_start.configure, state="normal")
            self.ui_call(self.btn_live.configure, state="normal")
            self.ui_call(self.live_engine_menu.configure, state="normal")
            self.ui_call(self.btn_lang.configure, state="normal")
            self.ui_call(self.btn_stop.configure, state="disabled", text=self.i18n[self.current_lang]["btn_stop"])
            self.msg_queue.put({"type": "progress", "val": 0})
            t = self.i18n[self.current_lang]
            kind, payload = _result
            if kind == "error":
                self.ui_call(messagebox.showerror, t["mb_title_error"], payload)
            elif kind == "cancelled":
                self.ui_call(messagebox.showwarning, t["mb_title_cancelled"], t["mb_msg_cancelled"])
            elif kind == "success":
                self.ui_call(messagebox.showinfo, t["mb_title_success"], t["mb_msg_success"] + f"\n\nFolder: {payload}")
                self.ui_call(self.open_folder, payload)

if __name__ == "__main__":
    app = PlatinumTranscriberApp()
    app.mainloop()